from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional

import pdfplumber
from docx import Document

from ..logging_config import get_logger


log = get_logger("service.parser")


@dataclass
class PageText:
    page_number: int
    text: str


@dataclass
class ParsedDocument:
    filename: str
    full_text: str
    pages: List[PageText]
    doc_type: str  # "pdf" | "docx" | "txt"
    is_scanned_suspected: bool
    text_density_per_page: List[float]
    ocr_used: bool = False  # NEW: Track if OCR was used


class ParserService:
    """
    Deterministic local parser for PDF and DOCX documents.
    Automatically uses OCR for scanned PDFs.
    """

    def __init__(self) -> None:
        self._ocr_service = None
    
    def _get_ocr(self):
        """Lazy load OCR service."""
        if self._ocr_service is None:
            from .ocr_service import get_ocr_service
            self._ocr_service = get_ocr_service()
        return self._ocr_service

    def parse(self, content: bytes, filename: str) -> ParsedDocument:
        suffix = Path(filename).suffix.lower()
        if suffix.endswith("pdf"):
            return self._parse_pdf_bytes(content, filename)
        if suffix.endswith("docx"):
            return self._parse_docx_bytes(content, filename)
        
        # Fallback: treat as plain text
        text = content.decode("utf-8", errors="ignore")
        pages = [PageText(page_number=1, text=text)]
        density = [self._compute_density(text)]
        return ParsedDocument(
            filename=filename,
            full_text=text,
            pages=pages,
            doc_type="txt",
            is_scanned_suspected=False,
            text_density_per_page=density,
            ocr_used=False,
        )

    def _parse_pdf_bytes(self, content: bytes, filename: str) -> ParsedDocument:
        from io import BytesIO

        pages: List[PageText] = []
        densities: List[float] = []
        full_text_chunks: List[str] = []
        
        with pdfplumber.open(BytesIO(content)) as pdf:
            for idx, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                pages.append(PageText(page_number=idx, text=text))
                full_text_chunks.append(text)
                densities.append(self._compute_density(text))
        
        full_text = "\n\n".join(full_text_chunks)
        
        # Detect if PDF is scanned (low text density)
        avg_density = sum(densities) / len(densities) if densities else 0
        is_scanned = avg_density < 0.1 or any(d < 0.05 for d in densities)
        
        # If scanned or very low text, return with flag for async OCR
        if is_scanned and len(full_text.strip()) < 100:
            log.warning(f"PDF appears to be scanned (avg density: {avg_density:.3f}). OCR recommended.")
        
        return ParsedDocument(
            filename=filename,
            full_text=full_text,
            pages=pages,
            doc_type="pdf",
            is_scanned_suspected=is_scanned,
            text_density_per_page=densities,
            ocr_used=False,
        )

    async def parse_with_ocr(self, content: bytes, filename: str) -> ParsedDocument:
        """
        Parse document, automatically using OCR for scanned PDFs.
        This is the preferred method for analysis.
        """
        # First, try normal parsing
        result = self.parse(content, filename)
        
        # If it's a PDF with low text content, use OCR
        if result.doc_type == "pdf":
            total_text_len = len(result.full_text.strip())
            avg_density = sum(result.text_density_per_page) / len(result.text_density_per_page) if result.text_density_per_page else 0
            
            # Use OCR if: very low density OR very little text extracted
            should_ocr = (avg_density < 0.1) or (total_text_len < 200 and result.is_scanned_suspected)
            
            if should_ocr:
                log.info(f"Using OCR for scanned PDF: {filename} (density: {avg_density:.3f}, text_len: {total_text_len})")
                
                ocr = self._get_ocr()
                if ocr.enabled:
                    ocr_text = await ocr.extract_text_from_pdf_bytes(content)
                    
                    if ocr_text and len(ocr_text) > len(result.full_text):
                        log.info(f"OCR successful: extracted {len(ocr_text)} chars (was {total_text_len})")
                        
                        # Parse OCR text into pages
                        ocr_pages = self._parse_ocr_text_to_pages(ocr_text)
                        
                        return ParsedDocument(
                            filename=filename,
                            full_text=ocr_text,
                            pages=ocr_pages,
                            doc_type="pdf",
                            is_scanned_suspected=True,
                            text_density_per_page=[self._compute_density(ocr_text)],
                            ocr_used=True,
                        )
                    else:
                        log.warning("OCR did not improve extraction")
                else:
                    log.warning("OCR service not available")
        
        return result
    
    def _parse_ocr_text_to_pages(self, ocr_text: str) -> List[PageText]:
        """
        Parse OCR text that contains [Page N] markers into PageText objects.
        """
        import re
        
        pages = []
        # Split by page markers
        parts = re.split(r'\[Page (\d+)\]', ocr_text)
        
        # parts will be: ['', '1', 'page1_text', '2', 'page2_text', ...]
        i = 1
        while i < len(parts) - 1:
            page_num = int(parts[i])
            page_text = parts[i + 1].strip()
            pages.append(PageText(page_number=page_num, text=page_text))
            i += 2
        
        # If no page markers found, treat as single page
        if not pages:
            pages.append(PageText(page_number=1, text=ocr_text))
        
        return pages

    def _parse_docx_bytes(self, content: bytes, filename: str) -> ParsedDocument:
        from io import BytesIO

        doc = Document(BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs]
        full_text = "\n".join(paragraphs)
        pages = [PageText(page_number=1, text=full_text)]
        density = [self._compute_density(full_text)]
        
        return ParsedDocument(
            filename=filename,
            full_text=full_text,
            pages=pages,
            doc_type="docx",
            is_scanned_suspected=False,
            text_density_per_page=density,
            ocr_used=False,
        )

    @staticmethod
    def _compute_density(text: str) -> float:
        if not text:
            return 0.0
        # simple heuristic: non-whitespace chars per total chars
        non_ws = sum(1 for c in text if not c.isspace())
        return non_ws / max(len(text), 1)


