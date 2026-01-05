"""
DOCX parser for extracting text from contract Word documents.
Uses python-docx for reliable extraction.
"""
from pathlib import Path
from typing import Optional
from docx import Document
from loguru import logger


class DOCXParser:
    """
    Extracts text from DOCX files while preserving structure.
    
    Design choice: python-docx is the standard library for .docx files.
    It preserves paragraph structure which is important for clause detection.
    """
    
    def __init__(self):
        pass
    
    def extract_text(self, file_path: Path) -> str:
        """
        Extract raw text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text as string
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file is not a valid DOCX
        """
        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
        if not file_path.suffix.lower() in [".docx", ".doc"]:
            raise ValueError(f"File is not a DOCX: {file_path}")
        
        logger.info(f"Extracting text from DOCX: {file_path.name}")
        
        try:
            doc = Document(file_path)
            
            # Extract text from all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # Skip empty paragraphs
                    paragraphs.append(text)
            
            # Extract text from tables (contracts often have tables)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)
            
            result = "\n\n".join(paragraphs)
            logger.info(f"Extracted {len(result)} total characters from DOCX")
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise ValueError(f"Failed to parse DOCX: {e}")
    
    def extract_metadata(self, file_path: Path) -> dict:
        """
        Extract basic metadata from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary with file metadata
        """
        metadata = {
            "filename": file_path.name,
            "file_size": file_path.stat().st_size,
            "file_type": "docx",
        }
        
        try:
            doc = Document(file_path)
            metadata["num_paragraphs"] = len(doc.paragraphs)
            metadata["num_tables"] = len(doc.tables)
        except Exception:
            metadata["num_paragraphs"] = 0
            metadata["num_tables"] = 0
        
        return metadata

